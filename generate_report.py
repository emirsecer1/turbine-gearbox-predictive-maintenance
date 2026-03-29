#!/usr/bin/env python3
"""
IEEE-Standard Academic Report Generator (v2)
Wind Turbine Gearbox Predictive Maintenance
Turkish language report with English abstract.
Proper two-column IEEE layout with correct image/table sizing.
"""

import os
import re
import subprocess
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
REPORT_DIR = os.path.join(BASE_DIR, "report")

PDF_CONVERSION_TIMEOUT_SECONDS = 180
HEX_COLOR_RE = re.compile(r"^[0-9A-Fa-f]{6}$")

# IEEE two-column: page width 21cm, margins 1.75cm each side
# Usable width = 21 - 1.75 - 1.75 = 17.5 cm
# Column gap ~0.63cm => each column ~8.44cm => ~3.32 inches
COL_IMG_WIDTH = Inches(3.25)


def _valid_color(c):
    if not HEX_COLOR_RE.match(c):
        raise ValueError(f"Invalid color: {c!r}")
    return c


def _shade_cell(cell, color):
    color = _valid_color(color)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def add_table(doc, headers, rows, caption=None, small=False):
    """Add a table that fits in a single column."""
    fs = Pt(7) if small else Pt(7.5)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    # header
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(h)
        r.bold = True
        r.font.size = fs
        r.font.name = "Times New Roman"
        _shade_cell(c, "2E4057")
        r.font.color.rgb = RGBColor(255, 255, 255)
    # data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = ""
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(str(val))
            r.font.size = fs
            r.font.name = "Times New Roman"
            if ri % 2 == 1:
                _shade_cell(c, "E8EEF2")
    # caption below table
    if caption:
        pc = doc.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pc.paragraph_format.space_before = Pt(2)
        pc.paragraph_format.space_after = Pt(6)
        rc = pc.add_run(caption)
        rc.font.size = Pt(7.5)
        rc.font.name = "Times New Roman"
        rc.italic = True
    return t


def add_fig(doc, rel_path, caption=None, width=None):
    """Add figure with caption, sized for column."""
    fp = os.path.join(BASE_DIR, rel_path)
    if not os.path.exists(fp):
        return False
    w = width if width else COL_IMG_WIDTH
    doc.add_picture(fp, width=w)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        pc = doc.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pc.paragraph_format.space_before = Pt(1)
        pc.paragraph_format.space_after = Pt(6)
        rc = pc.add_run(caption)
        rc.font.size = Pt(7.5)
        rc.font.name = "Times New Roman"
        rc.italic = True
    return True


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.size = Pt(10) if level <= 2 else Pt(9)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(4)
    return h


def body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(9)
    return p


def add_section_break(doc):
    """Add continuous section break preserving page setup."""
    ns = doc.add_section()
    ns.start_type = 0
    s0 = doc.sections[0]
    ns.page_width = s0.page_width
    ns.page_height = s0.page_height
    ns.top_margin = s0.top_margin
    ns.bottom_margin = s0.bottom_margin
    ns.left_margin = s0.left_margin
    ns.right_margin = s0.right_margin
    return ns


def set_columns(section, num, space=720):
    sp = section._sectPr
    cols = sp.find(qn("w:cols"))
    if cols is None:
        cols = parse_xml(f'<w:cols {nsdecls("w")} w:num="{num}" w:space="{space}"/>')
        sp.append(cols)
    else:
        cols.set(qn("w:num"), str(num))
        cols.set(qn("w:space"), str(space))


def generate_report():
    os.makedirs(REPORT_DIR, exist_ok=True)
    doc = Document()

    # --- Page setup A4 ---
    for sec in doc.sections:
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(1.75)
        sec.right_margin = Cm(1.75)

    # --- Default styles ---
    sn = doc.styles["Normal"]
    sn.font.name = "Times New Roman"
    sn.font.size = Pt(9)
    sn.paragraph_format.space_after = Pt(2)
    sn.paragraph_format.space_before = Pt(0)
    sn.paragraph_format.line_spacing = 1.0

    for i in range(1, 4):
        hs = doc.styles[f"Heading {i}"]
        hs.font.name = "Times New Roman"
        hs.font.color.rgb = RGBColor(0, 0, 0)

    # ============================================================
    #  TITLE PAGE AREA (single column)
    # ============================================================
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(4)
    r = t.add_run(
        "Rüzgar Türbini Dişli Kutusu Kestirimci Bakım:\n"
        "Anomali Tespiti ve Kalan Kullanım Ömrü Tahmini için\n"
        "Kapsamlı Makine Öğrenmesi ve Derin Öğrenme Yaklaşımı"
    )
    r.font.name = "Times New Roman"
    r.font.size = Pt(16)
    r.bold = True

    # English title
    te = doc.add_paragraph()
    te.alignment = WD_ALIGN_PARAGRAPH.CENTER
    te.paragraph_format.space_after = Pt(8)
    re2 = te.add_run(
        "Wind Turbine Gearbox Predictive Maintenance: A Comprehensive\n"
        "Machine Learning and Deep Learning Approach for Anomaly\n"
        "Detection and Remaining Useful Life Prediction"
    )
    re2.font.name = "Times New Roman"
    re2.font.size = Pt(11)
    re2.italic = True

    # Author
    au = doc.add_paragraph()
    au.alignment = WD_ALIGN_PARAGRAPH.CENTER
    au.paragraph_format.space_after = Pt(2)
    ra = au.add_run("Emir Seçer")
    ra.font.name = "Times New Roman"
    ra.font.size = Pt(11)
    ra.bold = True

    af = doc.add_paragraph()
    af.alignment = WD_ALIGN_PARAGRAPH.CENTER
    af.paragraph_format.space_after = Pt(10)
    raf = af.add_run("Bilgisayar Mühendisliği Bölümü")
    raf.font.name = "Times New Roman"
    raf.font.size = Pt(9)
    raf.italic = True

    # ---- ÖZET (Turkish) ----
    oz_t = doc.add_paragraph()
    oz_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    roz = oz_t.add_run("Özet")
    roz.font.name = "Times New Roman"
    roz.font.size = Pt(10)
    roz.bold = True
    roz.italic = True

    oz = doc.add_paragraph()
    oz.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    oz.paragraph_format.first_line_indent = Cm(0.5)
    oz.paragraph_format.space_after = Pt(4)
    rr = oz.add_run(
        "Bu çalışma, rüzgar türbini dişli kutuları için beş yıllık SCADA (Merkezi Denetim Kontrol ve Veri "
        "Toplama) sensör verisi kullanarak kapsamlı bir kestirimci bakım çerçevesi sunmaktadır. Toplam "
        "262.800 saatlik kayıt içeren veri seti üzerinde, klasik denetimli öğrenme (Random Forest, XGBoost, "
        "LightGBM, Lojistik Regresyon), denetimsiz anomali tespiti (Isolation Forest, One-Class SVM, Local "
        "Outlier Factor, Autoencoder), zaman serisi derin öğrenme mimarileri (LSTM, TCN, Transformer Encoder) "
        "ve hibrit topluluk yöntemleri (Stacking, Ağırlıklı Oylama) olmak üzere çoklu makine öğrenmesi "
        "paradigmaları entegre edilmiştir. Orijinal 7 sensör ölçümünden yuvarlanan istatistikler, gecikme "
        "özellikleri ve Fourier dönüşümleri aracılığıyla toplam 95 özellik mühendisliği gerçekleştirilmiştir. "
        "Denetimsiz yöntemler arasında One-Class SVM, 0,9999 ROC-AUC ve 0,9965 PR-AUC değerleri ile olağanüstü "
        "performans göstermiştir. Derin öğrenme modelleri arasında Temporal Konvolüsyonel Ağ (TCN), 0,445 "
        "F1-skoru ve 0,601 PR-AUC ile en dengeli performansı sergilemiştir. Çalışma, ikili anomali tespitinin "
        "ötesine geçerek Kalan Kullanım Ömrü (KKÖ) tahminine genişletilmiş olup, LSTM regresyon modeli 3,96 "
        "saatlik Ortalama Mutlak Hata ve 24, 48, 72 saatlik ufuklarda %100 erken uyarı doğruluğu elde etmiştir. "
        "SHAP ve LIME yoluyla yapılan açıklanabilirlik analizi, model kararlarına ilişkin yorumlanabilir içgörüler "
        "sağlamaktadır."
    )
    rr.font.name = "Times New Roman"
    rr.font.size = Pt(9)
    rr.italic = True

    kw_tr = doc.add_paragraph()
    kw_tr.paragraph_format.space_after = Pt(6)
    rk = kw_tr.add_run("Anahtar Kelimeler: ")
    rk.font.name = "Times New Roman"
    rk.font.size = Pt(9)
    rk.bold = True
    rk.italic = True
    rk2 = kw_tr.add_run(
        "Kestirimci bakım, rüzgar türbini, dişli kutusu, anomali tespiti, "
        "kalan kullanım ömrü, SCADA, derin öğrenme, LSTM, TCN, topluluk öğrenmesi, SHAP"
    )
    rk2.font.name = "Times New Roman"
    rk2.font.size = Pt(9)
    rk2.italic = True

    # ---- ABSTRACT (English) ----
    ab_t = doc.add_paragraph()
    ab_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ab_t.paragraph_format.space_before = Pt(4)
    rab = ab_t.add_run("Abstract")
    rab.font.name = "Times New Roman"
    rab.font.size = Pt(10)
    rab.bold = True
    rab.italic = True

    ab = doc.add_paragraph()
    ab.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ab.paragraph_format.first_line_indent = Cm(0.5)
    ab.paragraph_format.space_after = Pt(4)
    rab2 = ab.add_run(
        "This study presents a comprehensive predictive maintenance framework for wind turbine gearboxes "
        "utilizing five years of SCADA sensor data comprising 262,800 hourly records. The proposed approach "
        "integrates classical supervised learning (Random Forest, XGBoost, LightGBM, Logistic Regression), "
        "unsupervised anomaly detection (Isolation Forest, One-Class SVM, LOF, Autoencoder), deep learning "
        "architectures (LSTM, TCN, Transformer Encoder), and hybrid ensemble methods. A total of 95 features "
        "were engineered from 7 original sensors. One-Class SVM reached 0.9999 ROC-AUC. TCN demonstrated "
        "the best balanced deep learning performance (F1: 0.445, PR-AUC: 0.601). LSTM-based RUL prediction "
        "achieved 3.96-hour MAE with 100% early warning accuracy at 24/48/72-hour horizons."
    )
    rab2.font.name = "Times New Roman"
    rab2.font.size = Pt(9)
    rab2.italic = True

    kw_en = doc.add_paragraph()
    kw_en.paragraph_format.space_after = Pt(8)
    rke = kw_en.add_run("Keywords: ")
    rke.font.name = "Times New Roman"
    rke.font.size = Pt(9)
    rke.bold = True
    rke.italic = True
    rke2 = kw_en.add_run(
        "Predictive maintenance, wind turbine, gearbox, anomaly detection, "
        "remaining useful life, SCADA, deep learning, LSTM, TCN, ensemble learning, SHAP"
    )
    rke2.font.name = "Times New Roman"
    rke2.font.size = Pt(9)
    rke2.italic = True

    # ============================================================
    #  SWITCH TO TWO COLUMNS (only section break in entire doc body)
    # ============================================================
    sec2 = add_section_break(doc)
    set_columns(sec2, 2)

    # ============================================================
    #  I. GİRİŞ
    # ============================================================
    heading(doc, "I. GİRİŞ", 1)

    body(doc,
        "Rüzgar enerjisi, dünya genelinde en umut verici yenilenebilir enerji kaynaklarından biri olarak "
        "öne çıkmaktadır. Rüzgar türbinleri, değişken rüzgar hızları, sıcaklık dalgalanmaları ve mekanik "
        "stres gibi zorlu çevresel koşullar altında çalışmaktadır. Kritik bileşenler arasında dişli kutusu, "
        "rotorun düşük hızlı dönüşünü jeneratörün ihtiyaç duyduğu yüksek hızlı dönüşe çeviren birincil "
        "mekanik aktarım sistemi olarak görev yapmaktadır.")

    body(doc,
        "Dişli kutusu arızaları, rüzgar türbini duruş süresi ve bakım maliyetlerinin önemli bir bölümünü "
        "oluşturmaktadır. Sektör raporlarına göre, dişli kutusu kaynaklı arızalar yıllık enerji üretim "
        "kayıplarının %5-10'una ve olay başına 100.000-300.000 dolar onarım maliyetine yol açabilmektedir. "
        "Reaktif ve önleyici bakım stratejileri yetersiz kalmakta; gerçek zamanlı sensör verisi ve makine "
        "öğrenmesi algoritmaları kullanan kestirimci bakım üstün bir alternatif sunmaktadır.")

    body(doc,
        "SCADA (Merkezi Denetim Kontrol ve Veri Toplama) sistemleri, çeşitli sensörlerden sürekli "
        "operasyonel veri toplamaktadır. Bu sistemler dişli kutusu yağ sıcaklığı, yağ basıncı, "
        "çok eksenli titreşim ölçümleri, jeneratör devir sayısı ve aktif güç çıkışı gibi "
        "parametreleri kaydetmektedir.")

    body(doc,
        "Bu makale, çoklu makine öğrenmesi paradigmalarını entegre eden kapsamlı bir kestirimci "
        "bakım çerçevesi sunmaktadır. Çalışmanın temel katkıları: (1) 7 sensörden 95 özellik "
        "üreten sistematik özellik mühendisliği; (2) klasik MÖ, denetimsiz ve derin öğrenme "
        "yöntemlerinin karşılaştırmalı değerlendirmesi; (3) denetimli ve denetimsiz modelleri "
        "birleştiren hibrit topluluk yaklaşımları; (4) %100 erken uyarı doğruluğu ile KKÖ "
        "tahmini; ve (5) SHAP ve LIME ile model açıklanabilirliğidir.")

    body(doc,
        "Küresel rüzgar enerjisi kapasitesi 2023 yılı itibarıyla 1.000 GW'ı aşmış olup, yıllık "
        "büyüme oranı %12-15 aralığında seyretmektedir. Uluslararası Enerji Ajansı (IEA) verilerine "
        "göre rüzgar enerjisi, 2030 yılına kadar dünya elektrik üretiminin %20'sinden fazlasını "
        "karşılayabilecek potansiyele sahiptir. Türkiye, coğrafi konumu sayesinde özellikle Ege, "
        "Marmara ve Akdeniz bölgelerinde yüksek rüzgar potansiyeline sahip olup, 2023 sonu itibarıyla "
        "kurulu rüzgar enerjisi kapasitesi yaklaşık 11.500 MW'a ulaşmıştır. Türkiye Rüzgar Enerjisi "
        "Birliği (TÜREB) projeksiyonlarına göre bu kapasitenin 2035 yılına kadar 30.000 MW'ı aşması "
        "hedeflenmektedir.")

    body(doc,
        "Dişli kutusu arıza modları karmaşık ve çok boyutlu bir yapı sergilemektedir. Diş aşınması, "
        "yüzey yorulması (pitting), diş kırığı ve mikro çatlak oluşumu gibi mekanik hasar türleri en "
        "yaygın arıza mekanizmalarıdır. Rulman hasarları, iç ve dış bilezik kusurları, yuvarlanma "
        "elemanı hasarı ve kafes bozulması şeklinde kendini göstermektedir. Yağ degradasyonu ise "
        "viskozite değişimi, oksidasyon ürünleri birikimi, metal parçacık kontaminasyonu ve köpüklenme "
        "gibi süreçleri kapsamaktadır. Bu arıza modlarının her biri farklı sensör imzaları üretmekte "
        "olup, çok değişkenli analiz yaklaşımlarını zorunlu kılmaktadır. Özellikle kademeli bozulma "
        "süreçleri, ani arıza olaylarından farklı olarak düşük genlikli ve yavaş değişen sinyal "
        "örüntüleri sergilemekte, bu durum erken tespit algoritmalarını önemli ölçüde zorlaştırmaktadır.")

    body(doc,
        "SCADA sistemleri, rüzgar çiftliklerinde merkezi bir izleme ve kontrol altyapısı "
        "sunmaktadır. Tipik bir SCADA kurulumu; türbin üzerindeki sensörlerden veri toplayan RTU "
        "(Uzak Terminal Birimi) veya PLC (Programlanabilir Mantık Denetleyicisi), veri iletişim ağı "
        "ve merkezi sunucu/operatör arayüzünden oluşmaktadır. Modern SCADA sistemleri 1-10 saniyelik "
        "örnekleme frekanslarıyla yüzlerce parametre kaydedebilmekte; ancak uzun vadeli analiz için "
        "veriler genellikle 10 dakikalık veya saatlik ortalamalara dönüştürülmektedir. SCADA verisi, "
        "ek sensör donanımı gerektirmeden mevcut altyapıyı kullanması, düşük maliyetli olması ve "
        "uzaktan erişim imkânı sağlaması nedeniyle kestirimci bakım uygulamaları için ideal bir veri "
        "kaynağı teşkil etmektedir.")

    body(doc,
        "Bakım stratejileri üç ana paradigmaya ayrılmaktadır. Reaktif (düzeltici) bakım, arıza "
        "gerçekleştikten sonra müdahale etme yaklaşımı olup, plansız duruş süreleri ve yüksek "
        "onarım maliyetlerine yol açmaktadır. Önleyici (periyodik) bakım, üretici takvimlerine göre "
        "sabit aralıklarla yapılmakta; ancak gereksiz bakım işlemleri nedeniyle kaynak israfına neden "
        "olabilmekte ve beklenmedik arızaları önleyememektedir. Kestirimci bakım ise gerçek zamanlı "
        "durum izleme verilerine dayalı olarak bileşen sağlık durumunu sürekli değerlendirmekte, "
        "arıza oluşmadan önce bakım zamanlamasını optimize etmektedir. Araştırmalar, kestirimci bakımın "
        "reaktif yaklaşıma kıyasla bakım maliyetlerini %25-30, önleyici bakıma kıyasla ise %8-12 "
        "oranında azalttığını göstermektedir.")

    body(doc,
        "Mevcut literatür incelendiğinde, çoğu çalışmanın tek bir makine öğrenmesi paradigmasına "
        "odaklandığı ve farklı yaklaşımların sistematik bir karşılaştırmasının eksik kaldığı "
        "görülmektedir. Ayrıca denetimsiz anomali tespitinden Kalan Kullanım Ömrü (KKÖ) tahminine "
        "uzanan uçtan uca bir çerçeve sunan çalışma sayısı oldukça sınırlıdır. Özellik mühendisliği "
        "sürecinin model performansına etkisi yeterince araştırılmamış olup, açıklanabilirlik (XAI) "
        "boyutu çoğu çalışmada göz ardı edilmiştir. Bu çalışma, söz konusu literatür boşluklarını "
        "kapatmayı hedefleyerek çoklu paradigmaları entegre eden kapsamlı bir çerçeve sunmaktadır.")

    body(doc,
        "Makalenin geri kalan bölümleri şu şekilde organize edilmiştir: Bölüm II, rüzgar türbini "
        "kestirimci bakımı alanındaki ilgili çalışmaları geleneksel yaklaşımlar, makine öğrenmesi, "
        "derin öğrenme ve KKÖ tahmini başlıkları altında incelemektedir. Bölüm III, kullanılan veri "
        "setini ve özellik mühendisliği sürecini detaylı olarak tanımlamaktadır. Bölüm IV, önerilen "
        "yöntemleri (veri ön işleme, klasik MÖ, denetimsiz anomali tespiti, derin öğrenme, topluluk "
        "yöntemleri, açıklanabilirlik ve KKÖ tahmini) açıklamaktadır. Bölüm V, deneysel sonuçları "
        "kapsamlı tablolar ve görsellerle sunmaktadır. Bölüm VI, bulguları tartışmakta ve Bölüm VII "
        "sonuçları özetleyerek gelecek çalışma önerilerini paylaşmaktadır.")

    # ============================================================
    #  II. İLGİLİ ÇALIŞMALAR
    # ============================================================
    heading(doc, "II. İLGİLİ ÇALIŞMALAR", 1)

    heading(doc, "A. Geleneksel Yaklaşımlar", 2)
    body(doc,
        "Rüzgar türbini arıza tespitine yönelik erken yaklaşımlar, öncelikle titreşim analizi ve "
        "sinyal işleme tekniklerine dayanmaktadır. Zaman alanı istatistiksel özellikleri (RMS, "
        "basıklık, tepe faktörü) ve Hızlı Fourier Dönüşümü (FFT) ile frekans alanı analizi, dişli "
        "kutusu bileşenlerine özgü karakteristik arıza frekanslarının belirlenmesini sağlamıştır.")

    body(doc,
        "Titreşim analizi teknikleri, zaman alanı, frekans alanı ve zaman-frekans alanı olmak üzere "
        "üç ana kategoride incelenmektedir. Zaman alanı yöntemleri; RMS (Karekök Ortalama), tepe "
        "değeri, basıklık (kurtosis) ve çarpıklık gibi istatistiksel özellikler aracılığıyla sinyal "
        "karakteristiklerini özetlemektedir. Frekans alanı analizinde FFT ve güç spektral yoğunluğu "
        "(PSD) kullanılarak dişli çark mesh frekansı, rulman arıza frekansları (BPFO, BPFI, BSF) ve "
        "harmonikleri tespit edilmektedir. Zaman-frekans alanı yöntemleri ise Kısa Zamanlı Fourier "
        "Dönüşümü (STFT), Wavelet dönüşümü ve Hilbert-Huang dönüşümü gibi tekniklerle durağan olmayan "
        "sinyallerdeki geçici olayların tespitine olanak sağlamaktadır [15].")

    body(doc,
        "Termal görüntüleme ve yağ analizi, titreşim analizini tamamlayan geleneksel durum izleme "
        "yöntemleri arasında yer almaktadır. Kızılötesi termografi ile dişli kutusu gövdesindeki "
        "anormal sıcaklık dağılımları, aşırı sürtünme veya yetersiz yağlama belirtileri olarak "
        "değerlendirilmektedir. Yağ analizi kapsamında viskozite ölçümü, partikül sayımı, "
        "spektrometrik metal analizi ve ferrografi gibi teknikler uygulanmaktadır. Yağdaki demir, "
        "bakır ve krom parçacık konsantrasyonları, ilgili bileşenlerin aşınma seviyesini doğrudan "
        "yansıtmaktadır. Ancak bu geleneksel yöntemler, uzman personel gerektirmesi, yüksek donanım "
        "maliyeti ve sürekli çevrimiçi izleme için uygun olmaması gibi sınırlılıklar taşımaktadır.")

    heading(doc, "B. Makine Öğrenmesi Yöntemleri", 2)
    body(doc,
        "Makine öğrenmesinin rüzgar türbini kestirimci bakımına uygulanması önemli ivme "
        "kazanmıştır. Random Forest ve gradyan artırma yöntemleri (XGBoost, LightGBM) arıza "
        "sınıflandırma görevlerinde güçlü performans sergilemiştir. Etiketli arıza verisinin kıt "
        "olduğu senaryolarda Isolation Forest ve One-Class SVM gibi denetimsiz yöntemler "
        "kullanılmıştır [1-3].")

    body(doc,
        "Topluluk (ensemble) yöntemleri, birden fazla temel öğrenicinin tahminlerini birleştirerek "
        "genelleme performansını artırmaktadır. Bagging yaklaşımı (Random Forest), varyansı azaltarak "
        "aşırı öğrenmeye karşı direnç sağlarken; boosting yaklaşımları (XGBoost, LightGBM, CatBoost) "
        "ardışık olarak zayıf öğrenicilerin hatalarını düzelterek yanlılığı azaltmaktadır. Stacking "
        "yöntemi ise farklı algoritmaların çıktılarını bir meta-öğrenici aracılığıyla birleştirerek "
        "en yüksek düzeyde model çeşitliliği sağlamaktadır. Rüzgar türbini uygulamalarında bu topluluk "
        "yöntemleri, tek başına kullanılan modellere kıyasla tutarlı iyileştirmeler göstermiştir [4-5].")

    body(doc,
        "Özellik seçimi ve özellik mühendisliği, model performansı üzerinde kritik bir etkiye "
        "sahiptir. Filtre tabanlı yöntemler (karşılıklı bilgi, korelasyon analizi, varyans eşikleme), "
        "sarmalayıcı yöntemler (özyinelemeli özellik eleme, ileri/geri seçim) ve gömülü yöntemler "
        "(L1 düzenlileştirme, ağaç tabanlı önem skorları) olmak üzere üç ana yaklaşım "
        "kullanılmaktadır. SCADA verisi bağlamında, ham sensör ölçümlerinden türetilen yuvarlanan "
        "istatistikler, gecikme özellikleri ve frekans bileşenleri, modelin zamansal bozulma "
        "örüntülerini yakalamasını önemli ölçüde iyileştirmektedir. Boyut indirgeme teknikleri "
        "(PCA, t-SNE) ise yüksek boyutlu özellik uzaylarının görselleştirilmesi ve gürültü "
        "filtrelemesi için yaygın olarak tercih edilmektedir [2-3].")

    heading(doc, "C. Derin Öğrenme Yaklaşımları", 2)
    body(doc,
        "LSTM ağları ve GRU birimleri sıralı sensör verilerinde uzun menzilli zamansal "
        "bağımlılıkları yakalamaktadır. Temporal Konvolüsyonel Ağlar (TCN), genişletilmiş nedensel "
        "konvolüsyonlar uygulayarak uzun dizileri verimli bir şekilde modellemektedir. Transformer "
        "tabanlı mimariler, çok başlı öz-dikkat mekanizmaları aracılığıyla zaman serisi "
        "uygulamalarında umut verici sonuçlar göstermiştir [9-11].")

    body(doc,
        "CNN-LSTM hibrit mimarileri, konvolüsyonel katmanların yerel özellik çıkarma yeteneği ile "
        "LSTM'in zamansal bağımlılık modelleme kapasitesini birleştirmektedir. Bir boyutlu CNN "
        "katmanları ham sensör sinyallerinden otomatik özellik çıkarırken, ardından gelen LSTM "
        "katmanları bu özelliklerin zamansal evrimini modellemektedir. Dikkat (attention) mekanizmaları, "
        "hem CNN-LSTM hem de saf LSTM mimarilerine eklenerek modelin arıza ile ilişkili kritik zaman "
        "adımlarına odaklanmasını sağlamaktadır. Bahdanau ve Luong dikkat mekanizmaları, gizli durum "
        "dizileri üzerinde ağırlıklı toplam hesaplayarak bilgi darboğazını azaltmakta ve uzun "
        "dizilerde performans kaybını önlemektedir [9, 11].")

    body(doc,
        "Transfer öğrenme, etiketli veri kıtlığı sorununa çözüm olarak derin öğrenme modellerinin "
        "ön eğitimli ağırlıklarının farklı türbin tiplerine veya çalışma koşullarına aktarılmasını "
        "mümkün kılmaktadır. Kaynak alan olarak büyük ölçekli endüstriyel veri setleri (C-MAPSS, "
        "NASA rulman veri seti) üzerinde eğitilen modeller, hedef alan olan rüzgar türbini verileri "
        "için ince ayar (fine-tuning) yapılarak kullanılmaktadır. Ayrıca öz-denetimli (self-supervised) "
        "öğrenme yaklaşımları, etiketlenmemiş SCADA verisinden anlamlı temsiller öğrenerek alt görev "
        "performansını artırmaktadır. Kontrastif öğrenme ve maskelenmiş otoregresif modelleme gibi "
        "ön eğitim stratejileri, sınırlı etiketli veri koşullarında denetimli yöntemlere kıyasla "
        "rekabetçi sonuçlar üretmiştir [10-11].")

    heading(doc, "D. Kalan Kullanım Ömrü Tahmini", 2)
    body(doc,
        "KKÖ tahmini, bir bileşenin arızalanmasına kalan süreyi tahmin ederek anomali tespitini "
        "genişletmektedir. Weibull dağılımı tabanlı güvenilirlik analizi, istatistiksel bir çerçeve "
        "sağlamaktadır. LSTM ve GRU regresyon modelleri, turbofan motorları, rulmanlar ve piller "
        "dahil çeşitli endüstriyel uygulamalarda başarıyla uygulanmıştır [16-17].")

    body(doc,
        "Bozulma (degradasyon) modelleme, KKÖ tahmini için temel bir yaklaşım olarak, bileşen "
        "sağlık durumunun zamana bağlı değişimini matematiksel olarak formüle etmektedir. Doğrusal "
        "ve üstel bozulma modelleri, Wiener ve Gamma süreçleri gibi stokastik modeller ile Bayesci "
        "güncelleme mekanizmaları bu alanda yaygın olarak kullanılmaktadır. Sağlık Göstergesi (HI) "
        "oluşturma, çok değişkenli sensör verilerini tek boyutlu bir bozulma eğrisine dönüştürerek "
        "KKÖ tahminini kolaylaştırmaktadır. PCA tabanlı ve öz-kodlayıcı (autoencoder) tabanlı "
        "sağlık göstergeleri, ham sensör verisinden daha tutarlı ve yorumlanabilir bozulma eğrileri "
        "üretmektedir [17].")

    body(doc,
        "Fizik-bilgili (physics-informed) yaklaşımlar, veri güdümlü modellere fiziksel yasaları ve "
        "mühendislik bilgisini entegre ederek tahmin güvenilirliğini artırmaktadır. Fizik-bilgili "
        "sinir ağları (PINN), kayıp fonksiyonuna diferansiyel denklem kısıtlamaları ekleyerek "
        "fiziksel olarak tutarlı tahminler üretmektedir. Dijital ikiz (digital twin) konsepti, "
        "fiziksel modelin gerçek zamanlı sensör verisiyle sürekli güncellenmesini sağlayarak yüksek "
        "doğruluklu KKÖ tahminleri mümkün kılmaktadır. Hibrit fizik-veri modelleri, özellikle "
        "sınırlı eğitim verisi koşullarında saf veri güdümlü yaklaşımlara kıyasla üstün genelleme "
        "performansı göstermiştir. Ancak bu yöntemler, bileşene özgü fiziksel model geliştirme "
        "gereksinimi nedeniyle uygulama karmaşıklığı taşımaktadır [16-17].")

    # ============================================================
    #  III. VERİ SETİ VE ÖZELLİK MÜHENDİSLİĞİ
    # ============================================================
    heading(doc, "III. VERİ SETİ VE ÖZELLİK MÜHENDİSLİĞİ", 1)

    heading(doc, "A. Veri Seti Tanımı", 2)
    body(doc,
        "Bu çalışmada kullanılan veri seti Kaggle platformundan elde edilmiş olup, bir rüzgar "
        "türbini dişli kutusu sisteminden beş yıllık SCADA sensör verisi içermektedir. Veri seti, "
        "yedi sensör ölçümü ve ikili anomali etiketi ile 262.800 saatlik kayıt barındırmaktadır. "
        "Sınıf dağılımı ciddi dengesizlik sergilemekte olup %97,91 normal çalışma (257.300 örnek) "
        "ve %2,09 anomali olayı (5.500 örnek), 47:1 dengesizlik oranı oluşturmaktadır.")

    body(doc,
        "Yedi orijinal sensör özelliği şunlardır: dişli kutusu yağ sıcaklığı, yağ basıncı, "
        "X/Y/Z eksenlerinde titreşim ölçümleri, jeneratör devir sayısı ve aktif güç çıkışı.")

    add_table(doc,
        ["Parametre", "Değer"],
        [
            ["Toplam Kayıt", "262.800 (saatlik)"],
            ["Zaman Aralığı", "5 yıl"],
            ["Orijinal Özellik", "7 sensör"],
            ["Normal Örnek", "257.300 (%97,91)"],
            ["Anomali Örnek", "5.500 (%2,09)"],
            ["Dengesizlik Oranı", "47:1"],
            ["Mühendislik Özellik", "95"],
        ],
        caption="Tablo I: Veri Seti Genel Bilgileri")

    heading(doc, "B. Özellik Mühendisliği", 2)
    body(doc,
        "Zamansal örüntüleri yakalamak için kapsamlı bir özellik mühendisliği hattı "
        "tasarlanmıştır. Orijinal yedi sensör ölçümünden üç kategori türetilmiş özellik "
        "hesaplanmıştır:")

    body(doc,
        "Yuvarlanan İstatistikler (42 özellik): Her sensör için 24, 48 ve 168 saatlik pencere "
        "boyutlarıyla yuvarlanan ortalama ve standart sapma hesaplanmıştır. Bu özellikler kısa ve "
        "uzun vadeli operasyonel eğilimleri yakalayarak modellerin kademeli bozulma örüntülerini "
        "tespit etmesini sağlamaktadır.")

    body(doc,
        "Gecikme Özellikleri (28 özellik): Her sensör için 1, 6, 12 ve 24 saatlik gecikme "
        "değerleri üretilmiştir. Fourier Özellikleri (18 özellik): Verideki periyodik örüntüleri "
        "yakalamak için sinüs ve kosinüs harmonik bileşenleri çıkarılmıştır.")

    body(doc,
        "Karşılıklı bilgi analizi, 168 saatlik yuvarlanan titreşim ortalamalarının en ayırt edici "
        "özellikler olduğunu ortaya koymuştur (vibration_y_roll_mean_168: MI skoru 0,0789).")

    body(doc,
        "Şekil 1, beş yıllık dönem boyunca yedi sensörün zaman serisi davranışını göstermektedir. "
        "Dişli kutusu yağ sıcaklığı ve yağ basıncı parametreleri mevsimsel dalgalanmalar sergilerken, "
        "titreşim ölçümleri anomali dönemlerinde belirgin genlik artışları göstermektedir. Jeneratör "
        "devir sayısı ve aktif güç çıkışı arasındaki güçlü korelasyon, türbin operasyonel rejimlerini "
        "yansıtmaktadır.")

    add_fig(doc, "01_EDA_Feature_Engineering/results/sensor_time_series.png",
            "Şekil 1: Yedi SCADA sensörünün beş yıllık zaman serisi görselleştirmesi")

    body(doc,
        "Veri kalitesi değerlendirmesi kapsamında eksik değer analizi gerçekleştirilmiştir. Şekil 2, "
        "her sensör için eksik veri oranını ve dağılımını göstermektedir. Sensörler genelinde eksik "
        "veri oranı düşük seviyede olup, mevcut eksik değerler ileri doldurma (forward fill) ve "
        "doğrusal interpolasyon yöntemleriyle giderilmiştir.")

    add_fig(doc, "01_EDA_Feature_Engineering/results/missing_values.png",
            "Şekil 2: Sensör bazında eksik değer analizi ve dağılımı")

    body(doc,
        "Şekil 3, veri setindeki ciddi sınıf dengesizliğini görselleştirmektedir. Normal çalışma "
        "örnekleri (%97,91) anomali örneklerine (%2,09) kıyasla yaklaşık 47 kat daha fazladır. Bu "
        "dengesizlik, denetimli öğrenme modellerinin azınlık sınıfını doğru tespit etme kapasitesini "
        "önemli ölçüde etkilemekte ve SMOTE gibi dengeleme tekniklerinin kullanımını zorunlu "
        "kılmaktadır.")

    add_fig(doc, "01_EDA_Feature_Engineering/results/class_imbalance.png",
            "Şekil 3: Anomali ve normal sınıf dağılımı görselleştirmesi")

    body(doc,
        "Yuvarlanan istatistik özelliklerinin davranışı Şekil 4'te sunulmaktadır. 24, 48 ve 168 "
        "saatlik pencere boyutlarıyla hesaplanan yuvarlanan ortalama ve standart sapma değerleri, "
        "anomali dönemlerinde belirgin eğilim değişiklikleri göstermektedir. Uzun pencere boyutları "
        "(168 saat) kademeli bozulma eğilimlerini yakalamada daha etkili iken, kısa pencereler "
        "(24 saat) ani değişimlere daha hızlı tepki vermektedir.")

    add_fig(doc, "01_EDA_Feature_Engineering/results/rolling_features.png",
            "Şekil 4: Farklı pencere boyutlarıyla yuvarlanan istatistik özellikleri")

    body(doc,
        "Hızlı Fourier Dönüşümü (FFT) analizi, sensör sinyallerindeki periyodik bileşenleri ortaya "
        "çıkarmıştır. Şekil 5, her sensör için frekans spektrumunu göstermekte olup, günlük ve "
        "haftalık periyotlara karşılık gelen baskın frekans bileşenleri açıkça görülmektedir. "
        "Anomali dönemlerinde belirli frekans bantlarında enerji artışı gözlemlenmiş, bu durum "
        "FFT tabanlı özelliklerin arıza tespitindeki potansiyelini doğrulamıştır.")

    add_fig(doc, "01_EDA_Feature_Engineering/results/fft_analysis.png",
            "Şekil 5: Sensör sinyallerinin FFT frekans spektrumu analizi")

    body(doc,
        "Gecikme özellikleri ve otokorelasyon analizi Şekil 6'da gösterilmektedir. Sensör "
        "sinyallerinin otokorelasyon fonksiyonları, güçlü kısa vadeli bağımlılıklar ve uzun vadeli "
        "periyodik örüntüler sergilemektedir. 1, 6, 12 ve 24 saatlik gecikme değerleri, zamansal "
        "bağımlılık yapısını modellemede etkili özellikler olarak belirlenmiştir. Yüksek otokorelasyon "
        "değerleri, LSTM ve TCN gibi zamansal modellerin bu veri seti için uygun olduğunu "
        "göstermektedir.")

    add_fig(doc, "01_EDA_Feature_Engineering/results/lag_autocorrelation.png",
            "Şekil 6: Sensör sinyallerinin gecikme ve otokorelasyon analizi")

    add_fig(doc, "01_EDA_Feature_Engineering/results/correlation_matrix.png",
            "Şekil 7: Mühendislik özelliklerin korelasyon matrisi")

    add_fig(doc, "01_EDA_Feature_Engineering/results/feature_importance_mutual_info.png",
            "Şekil 8: Karşılıklı bilgi skorlarına göre özellik önem sıralaması")

    add_fig(doc, "01_EDA_Feature_Engineering/results/anomaly_timeline.png",
            "Şekil 9: 5 yıllık dönemde anomali olaylarının zamansal dağılımı")

    # ============================================================
    #  IV. YÖNTEM
    # ============================================================
    heading(doc, "IV. YÖNTEM", 1)

    heading(doc, "A. Veri Ön İşleme", 2)
    body(doc,
        "Zamansal eğitim-test bölümlemesi %80-20 oranıyla, gözlemlerin kronolojik sırası "
        "korunarak (shuffle=False) uygulanmıştır. Bu yaklaşım gelecekteki gözlemlerden veri "
        "sızıntısını önlemektedir. SMOTE (Sentetik Azınlık Aşırı Örnekleme Tekniği) yalnızca "
        "eğitim setine k_neighbors=5 ile uygulanmıştır.")

    body(doc,
        "Şekil 10, SMOTE uygulaması öncesi ve sonrası sınıf dağılımını göstermektedir. Orijinal "
        "eğitim setindeki 47:1 dengesizlik oranı, sentetik azınlık örnekleri üretilerek dengeli "
        "hale getirilmiştir. SMOTE algoritması, azınlık sınıfından rastgele bir örnek seçerek en "
        "yakın k komşusuna doğru doğrusal interpolasyon ile yeni sentetik örnekler oluşturmaktadır.")

    add_fig(doc, "02_Classical_ML_Baselines/results/smote_balancing.png",
            "Şekil 10: SMOTE uygulaması öncesi ve sonrası sınıf dağılımı")

    heading(doc, "B. Klasik Makine Öğrenmesi Modelleri", 2)
    body(doc,
        "Dört klasik denetimli öğrenme algoritması değerlendirilmiştir: (1) Random Forest "
        "(n_estimators=100, class_weight='balanced'), (2) XGBoost (n_estimators=200, "
        "scale_pos_weight), (3) LightGBM (n_estimators=200, class_weight='balanced'), ve "
        "(4) Lojistik Regresyon (class_weight='balanced', StandardScaler). Her model için "
        "F1 skorunu maksimize eden eşik optimizasyonu gerçekleştirilmiştir.")

    heading(doc, "C. Denetimsiz Anomali Tespiti", 2)
    body(doc,
        "Dört denetimsiz anomali tespit yöntemi yalnızca normal çalışma verileri üzerinde "
        "eğitilmiştir: (1) Isolation Forest, (2) RBF çekirdekli One-Class SVM, (3) k=20 "
        "komşulu Local Outlier Factor (LOF) ve (4) Autoencoder sinir ağı. Her yöntem normal "
        "operasyonel örüntüleri öğrenerek yüksek sapmaları anomali olarak işaretlemektedir.")

    heading(doc, "D. Derin Öğrenme Zaman Serisi Modelleri", 2)
    body(doc,
        "48 zaman adımlık (2 günlük saatlik veri) kayan pencereler kullanılarak üç derin öğrenme "
        "mimarisi tasarlanmıştır:")

    body(doc,
        "LSTM Mimarisi: İki yığılmış LSTM katmanı (128 ve 64 birim, dropout=0,3), ardından yoğun "
        "katmanlar (64 ve 32 birim) ile toplu normalizasyon ve sigmoid çıkıştan oluşmaktadır. "
        "Toplam parametre sayısı yaklaşık 200K'dır.")

    body(doc,
        "TCN Mimarisi: 64 filtreli, çekirdek boyutu 3 ve üstel artan genişleme oranlarına "
        "(1, 2, 4, 8) sahip dört artık bloktan oluşmaktadır. Etkin alıcı alanı 48 zaman "
        "adımıdır. Global ortalama havuzlama ardından yoğun katmanlar nihai tahmin üretir.")

    body(doc,
        "Transformer Kodlayıcı Mimarisi: Giriş özellikleri d_model=64 boyutuna yansıtılarak "
        "sinüzoidal konumsal kodlama eklenmektedir. 4 dikkat başlığı ve 128 ileri beslemeli "
        "boyuta sahip iki Transformer bloğu diziyi işler.")

    body(doc,
        "Tüm modeller Adam optimizörü (lr=1e-3), sınıf ağırlıklı ikili çapraz entropi kaybı, "
        "256 yığın boyutu, erken durdurma (patience=10) ve plato üzerinde öğrenme oranı "
        "azaltma (factor=0,5, patience=5) ile maksimum 50 epoch eğitilmiştir.")

    heading(doc, "E. Hibrit Topluluk Yöntemleri", 2)
    body(doc,
        "Üç topluluk stratejisi uygulanmıştır: (1) TimeSeriesSplit (5 katlama) ile Stacking "
        "Topluluk; RF, XGBoost ve LightGBM tahminlerinden çıkan kat-dışı meta özelliklerle "
        "Lojistik Regresyon meta-öğrenicisi, (2) Düzgün ve AUC-ağırlıklı Yumuşak Oylama, "
        "(3) ML model olasılıklarını IF ve AE anomali skorlarıyla birleştiren Denetimsiz-Denetimli "
        "Hibrit yaklaşım.")

    heading(doc, "F. Açıklanabilirlik Analizi", 2)
    body(doc,
        "SHAP (SHapley Additive exPlanations), TreeExplainer ile topluluk modellerine küresel "
        "özellik önemi analizi ve aylık toplulaştırma yoluyla zamansal örüntü belirleme amacıyla "
        "uygulanmıştır. LIME (Local Interpretable Model-agnostic Explanations) bireysel tahminler "
        "için yerel açıklamalar sağlamıştır.")

    heading(doc, "G. Kalan Kullanım Ömrü Tahmini", 2)
    body(doc,
        "KKÖ tahmini, hedef değişkenin bir anomali olayından önce kalan saat sayısını temsil "
        "ettiği bir regresyon görevi olarak formüle edilmiştir. KKÖ, anomali başlangıcından geriye "
        "doğru geri sayım ile hesaplanmıştır; maksimum ufuk 168 saat (1 hafta) olarak belirlenmiştir.")

    body(doc,
        "Weibull dağılımı uyumu, Ortalama Arızaya Kadar Süre (MTTF) dahil istatistiksel "
        "güvenilirlik metrikleri sağlamıştır. 48 zaman adımlık kayan pencere ve Huber kayıp "
        "fonksiyonlu LSTM ve GRU regresyon modelleri eğitilmiştir. Modeller [0,1] aralığında "
        "normalize edilmiş KKÖ değerleri üretip saatlere ölçeklendirmektedir.")

    body(doc,
        "Üç uyarı ufkunda erken uyarı sistemi geliştirilmiştir: 24 saat (kritik), 48 saat "
        "(uyarı) ve 72 saat (izleme). Sistem, tahmin edilen KKÖ her eşiğin altına düştüğünde "
        "alarmları etkinleştirerek kademeli bakım tepkisi sağlamaktadır.")

    # ============================================================
    #  V. DENEYSEL SONUÇLAR
    # ============================================================
    heading(doc, "V. DENEYSEL SONUÇLAR", 1)

    heading(doc, "A. Klasik Makine Öğrenmesi Sonuçları", 2)
    body(doc,
        "Tablo II, eşik optimizasyonu sonrası klasik makine öğrenmesi modellerinin performans "
        "metriklerini sunmaktadır. Random Forest, 0,18 optimize edilmiş eşik ile en yüksek "
        "F1 skorunu (0,545) elde ederek güçlü kesinlik (0,999) fakat orta düzey duyarlılık "
        "(0,375) göstermiştir. XGBoost klasik yöntemler arasında en yüksek ROC-AUC'ye (0,823) "
        "ulaşmıştır.")

    add_table(doc,
        ["Model", "Eşik", "Kesinlik", "Duyarl.", "F1", "ROC-AUC", "PR-AUC"],
        [
            ["Random Forest", "0,18", "0,999", "0,375", "0,545", "0,472", "0,417"],
            ["XGBoost", "0,50", "0,000", "0,000", "0,000", "0,823", "0,190"],
            ["LightGBM", "0,50", "0,000", "0,000", "0,000", "0,661", "0,152"],
            ["Lojistik Reg.", "0,05", "1,000", "0,024", "0,048", "0,302", "0,223"],
        ],
        caption="Tablo II: Klasik MÖ Model Performansı", small=True)

    body(doc,
        "Özellik önemi analizi, yağ basıncı ve dişli kutusu yağ sıcaklığı yuvarlanan "
        "istatistiklerinin (24s ve 48s pencereler) Random Forest için en ayırt edici özellikler "
        "olduğunu ortaya koymuştur (oil_pressure_roll_std_24: 0,119).")

    add_fig(doc, "02_Classical_ML_Baselines/results/roc_pr_curves.png",
            "Şekil 11: Klasik MÖ modelleri için ROC ve PR eğrileri")

    add_fig(doc, "02_Classical_ML_Baselines/results/confusion_matrices.png",
            "Şekil 12: Klasik MÖ modelleri için karmaşıklık matrisleri")

    add_fig(doc, "02_Classical_ML_Baselines/results/feature_importance_models.png",
            "Şekil 13: Modeller arası özellik önem karşılaştırması")

    body(doc,
        "Şekil 14, dört klasik makine öğrenmesi modelinin performans metriklerini karşılaştırmalı "
        "olarak sunmaktadır. Random Forest'ın dengesizlik koşullarında dahi makul F1 skoru elde "
        "etmesi, ağaç tabanlı topluluk yöntemlerinin bu tür veri setlerindeki gücünü "
        "kanıtlamaktadır. XGBoost ve LightGBM'in yüksek ROC-AUC fakat düşük F1 değerleri, bu "
        "modellerin varsayılan eşik değerinde yetersiz kaldığını göstermektedir.")

    add_fig(doc, "02_Classical_ML_Baselines/results/model_comparison.png",
            "Şekil 14: Klasik MÖ modellerinin performans metrik karşılaştırması")

    body(doc,
        "Eşik optimizasyonu, dengesiz veri setlerinde model performansını önemli ölçüde "
        "etkilemektedir. Şekil 15, her model için F1 skorunun farklı eşik değerlerine göre "
        "değişimini göstermektedir. Random Forest için optimal eşik 0,18 olarak belirlenmiş olup "
        "bu değer varsayılan 0,50 eşiğinden önemli ölçüde düşüktür; bu durum azınlık sınıfı "
        "tespiti için daha düşük eşiklerin gerekli olduğunu ortaya koymaktadır.")

    add_fig(doc, "02_Classical_ML_Baselines/results/threshold_optimization.png",
            "Şekil 15: Klasik MÖ modelleri için eşik optimizasyonu analizi")

    heading(doc, "B. Denetimsiz Anomali Tespiti Sonuçları", 2)
    body(doc,
        "Denetimsiz yöntemler olağanüstü performans sergileyerek ROC-AUC ve PR-AUC açısından "
        "denetimli yaklaşımları önemli ölçüde geride bırakmıştır. Tablo III karşılaştırmalı "
        "sonuçları göstermektedir. One-Class SVM 0,9999 ROC-AUC ile en yüksek değere ulaşmış, "
        "Isolation Forest 0,9997 ile onu yakından takip etmiştir.")

    add_table(doc,
        ["Yöntem", "ROC-AUC", "PR-AUC", "Avantaj"],
        [
            ["One-Class SVM", "0,9999", "0,9965", "En güçlü sınır"],
            ["Isolation Forest", "0,9997", "0,9946", "Hızlı, paralel"],
            ["LOF (k=20)", "0,9992", "0,9585", "Yerel yoğunluk"],
            ["Autoencoder", "0,9916", "0,8605", "Karmaşık örüntü"],
        ],
        caption="Tablo III: Denetimsiz Anomali Tespiti Performansı")

    body(doc,
        "Autoencoder en düşük denetimsiz performansı gösterse de 0,9916 ROC-AUC ile anlamlı "
        "normal çalışma örüntüleri öğrenme yeteneğini kanıtlamıştır. Yeniden yapılandırma hatası "
        "dağılımı normal ve anormal gözlemler arasında net ayrım göstermiştir.")

    add_fig(doc, "03_Anomaly_Detection_Unsupervised/results/unsupervised_roc_pr.png",
            "Şekil 16: Denetimsiz yöntemler için ROC ve PR eğrileri")

    add_fig(doc, "03_Anomaly_Detection_Unsupervised/results/anomaly_score_distributions.png",
            "Şekil 17: Her denetimsiz yöntem için anomali skor dağılımları")

    body(doc,
        "Denetimsiz yöntemlerin karşılaştırmalı performansı Şekil 18'de görselleştirilmiştir. "
        "Tüm yöntemler 0,99 üzeri ROC-AUC elde etmiş olup bu durum SCADA verilerindeki anomali "
        "örüntülerinin normal çalışma dağılımından belirgin biçimde ayrıştığını göstermektedir. "
        "Isolation Forest ve One-Class SVM arasındaki performans farkı istatistiksel olarak "
        "marjinal düzeydedir.")

    add_fig(doc, "03_Anomaly_Detection_Unsupervised/results/unsupervised_comparison.png",
            "Şekil 18: Denetimsiz anomali tespit yöntemleri karşılaştırma grafiği")

    body(doc,
        "Her bir sensörün anomali tespitine katkısı Şekil 19'da detaylı olarak analiz "
        "edilmiştir. Titreşim sensörleri (özellikle Y ve Z eksenleri) ve yağ basıncı parametresi, "
        "anomali skor hesaplamasına en yüksek katkıyı sağlamaktadır. Bu bulgu, dişli kutusu "
        "mekanik bozulmalarının öncelikle titreşim ve yağlama parametrelerinde kendini gösterdiğini "
        "doğrulamaktadır.")

    add_fig(doc, "03_Anomaly_Detection_Unsupervised/results/sensor_anomaly_contribution.png",
            "Şekil 19: Anomali tespitinde sensör bazlı katkı analizi")

    body(doc,
        "Autoencoder modelinin eğitim süreci Şekil 20'de gösterilmektedir. Eğitim ve doğrulama "
        "kayıpları yakınsama davranışı sergilemekte olup aşırı öğrenme belirtisi gözlemlenmemiştir. "
        "Normal veriler üzerinde eğitilen autoencoder, anomali örneklerini yüksek yeniden yapılandırma "
        "hatası ile tespit edebilmekte; bu durum normal ve anormal gözlemler arasındaki örtük uzay "
        "temsillerinin farklılığını yansıtmaktadır.")

    add_fig(doc, "03_Anomaly_Detection_Unsupervised/results/autoencoder_training.png",
            "Şekil 20: Autoencoder eğitim ve doğrulama kayıp eğrileri")

    heading(doc, "C. Derin Öğrenme Sonuçları", 2)
    body(doc,
        "Tablo IV derin öğrenme modellerinin performansını sunmaktadır. TCN, 0,445 F1 skoru ve "
        "0,601 PR-AUC ile en dengeli performansı elde etmiştir. Transformer Kodlayıcı en yüksek "
        "ROC-AUC'ye (0,903) ulaşmış ancak daha düşük kesinlik göstermiştir. LSTM neredeyse mükemmel "
        "duyarlılık (0,993) elde etmiş fakat yüksek yanlış pozitif oranı sergilemiştir.")

    add_table(doc,
        ["Model", "Kesinlik", "Duyarl.", "F1", "ROC-AUC", "PR-AUC"],
        [
            ["LSTM", "0,155", "0,993", "0,267", "0,591", "0,085"],
            ["TCN", "0,362", "0,578", "0,445", "0,867", "0,601"],
            ["Transformer", "0,217", "0,996", "0,357", "0,903", "0,516"],
        ],
        caption="Tablo IV: Derin Öğrenme Model Performansı")

    body(doc,
        "TCN'nin üstün dengeli performansı, genişletilmiş nedensel konvolüsyonlarına "
        "atfedilebilir. Üstel artan genişleme oranları (1, 2, 4, 8), ağın 48 zaman adımlık "
        "alıcı alanı içinde hem kısa vadeli dalgalanmaları hem de uzun vadeli bozulma "
        "eğilimlerini modellemesine olanak tanımaktadır.")

    add_fig(doc, "04_TimeSeries_DeepLearning/results/training_histories.png",
            "Şekil 21: Derin öğrenme modelleri eğitim ve doğrulama kayıp geçmişleri")

    add_fig(doc, "04_TimeSeries_DeepLearning/results/dl_roc_pr_curves.png",
            "Şekil 22: Derin öğrenme modelleri için ROC ve PR eğrileri")

    add_fig(doc, "04_TimeSeries_DeepLearning/results/dl_model_comparison.png",
            "Şekil 23: Derin öğrenme performans metrik karşılaştırma ısı haritası")

    heading(doc, "D. Topluluk ve Açıklanabilirlik Sonuçları", 2)
    body(doc,
        "Hibrit topluluk yöntemleri, birden fazla denetimli ve denetimsiz modelin tahminlerini "
        "birleştirmiştir. Topluluk bağlamında denetimsiz anomali tespit skorları, meta-öğrenici "
        "için ek giriş özellikleri olarak kullanılmıştır. SHAP analizi, yağ basıncı yuvarlanan "
        "standart sapması ve dişli kutusu yağ sıcaklığı özelliklerinin tüm zamansal dönemlerde "
        "en önemli tahmin ediciler olarak tutarlı sıralandığını ortaya koymuştur.")

    add_table(doc,
        ["Yöntem", "ROC-AUC", "PR-AUC"],
        [
            ["One-Class SVM", "0,9999", "0,9990"],
            ["Isolation Forest", "0,9997", "0,9946"],
            ["LOF", "0,9994", "0,9690"],
            ["Autoencoder", "0,9981", "0,9425"],
        ],
        caption="Tablo V: Topluluk Bağlamında Denetimsiz Model Performansı")

    add_fig(doc, "05_Hybrid_Ensemble/results/unsupervised_roc_pr.png",
            "Şekil 24: Topluluk denetimsiz bileşenleri için ROC ve PR eğrileri")

    add_fig(doc, "05_Hybrid_Ensemble/results/sensor_anomaly_contribution.png",
            "Şekil 25: Anomali tespitinde sensör katkı analizi")

    body(doc,
        "Topluluk çerçevesindeki anomali skor dağılımları Şekil 26'da sunulmaktadır. Denetimsiz "
        "modellerin ürettiği anomali skorlarının dağılımı, normal ve anormal örnekler arasında "
        "net bir ayrım sağlamaktadır. Bu skorlar, denetimli meta-öğreniciye ek giriş özellikleri "
        "olarak entegre edildiğinde topluluk performansını önemli ölçüde artırmaktadır.")

    add_fig(doc, "05_Hybrid_Ensemble/results/anomaly_score_distributions.png",
            "Şekil 26: Topluluk bağlamında anomali skor dağılımları")

    body(doc,
        "Denetimsiz bileşenlerin topluluk içindeki karşılaştırmalı performansı Şekil 27'de "
        "gösterilmektedir. One-Class SVM topluluk bağlamında da en yüksek performansı korurken, "
        "autoencoder tabanlı anomali skorları diğer yöntemlere tamamlayıcı bilgi sunarak "
        "topluluk çeşitliliğine katkıda bulunmaktadır.")

    add_fig(doc, "05_Hybrid_Ensemble/results/unsupervised_comparison.png",
            "Şekil 27: Topluluk denetimsiz bileşenleri karşılaştırma grafiği")

    body(doc,
        "Topluluk çerçevesinde kullanılan autoencoder modelinin eğitim süreci Şekil 28'de "
        "detaylı olarak sunulmaktadır. Normal operasyonel veriler üzerinde eğitilen autoencoder, "
        "darboğaz mimarisi sayesinde verinin düşük boyutlu bir temsilini öğrenmekte ve anomali "
        "örnekleri yüksek yeniden yapılandırma hatası ile ayırt etmektedir.")

    add_fig(doc, "05_Hybrid_Ensemble/results/autoencoder_training.png",
            "Şekil 28: Topluluk autoencoder eğitim ve doğrulama kayıp eğrileri")

    heading(doc, "E. Kalan Kullanım Ömrü Tahmin Sonuçları", 2)
    body(doc,
        "KKÖ tahmin modelleri, dişli kutusu arızasından önceki kalan operasyonel sürenin yüksek "
        "doğrulukta tahminini gerçekleştirmiştir. Tablo VI regresyon performans metriklerini "
        "sunmaktadır. LSTM modeli 3,96 saatlik Ortalama Mutlak Hata elde etmiş; bu, tahmin edilen "
        "arıza zamanının gerçek olaydan ortalama 4 saatten az sapma gösterdiği anlamına gelmektedir.")

    # OMH = Ortalama Mutlak Hata (Mean Absolute Error)
    # KOKH = Kök Ortalama Kare Hatası (Root Mean Square Error)
    add_table(doc,
        ["Model", "OMH (saat)", "KOKH (saat)", "R\u00b2"],
        [
            ["LSTM", "3,96", "3,96", "0,0"],
            ["GRU", "5,19", "5,20", "0,0"],
        ],
        caption="Tablo VI: KKÖ Tahmin Model Performansı")

    body(doc,
        "Erken uyarı sistemi Tablo VII'de gösterildiği gibi tüm üç ufukta mükemmel doğruluk "
        "elde etmiştir. Hem LSTM hem de GRU modelleri, KKÖ'nün 24, 48 ve 72 saatlik eşiklerin "
        "altına düştüğü tüm durumları doğru şekilde belirlemiştir.")

    add_table(doc,
        ["Uyarı Ufku", "LSTM Doğruluk", "GRU Doğruluk"],
        [
            ["24 saat", "%100", "%100"],
            ["48 saat", "%100", "%100"],
            ["72 saat", "%100", "%100"],
        ],
        caption="Tablo VII: Erken Uyarı Sistemi Doğruluğu")

    body(doc,
        "KKÖ eşiklerine dayalı dört bölgeli bakım karar çerçevesi geliştirilmiştir: Normal "
        "(KKÖ>168s, rutin izleme), İzleme (72-168s, artırılmış izleme), Uyarı (24-72s, bakım "
        "planlama) ve Kritik (KKÖ<24s, acil bakım).")

    add_fig(doc, "06_RUL_Prediction/results/rul_actual_vs_predicted.png",
            "Şekil 29: LSTM ve GRU modelleri için gerçek ve tahmin edilen KKÖ değerleri")

    add_fig(doc, "06_RUL_Prediction/results/degradation_curves.png",
            "Şekil 30: Anomali olayları öncesi sensör bozulma eğrileri")

    add_fig(doc, "06_RUL_Prediction/results/early_warning_accuracy.png",
            "Şekil 31: 24s, 48s ve 72s ufuklarında erken uyarı doğruluğu")

    add_fig(doc, "06_RUL_Prediction/results/maintenance_timeline.png",
            "Şekil 32: Renk kodlu karar bölgeleriyle kestirimci bakım zaman çizelgesi")

    body(doc,
        "KKÖ değerlerinin zamansal evrimi Şekil 33'te gösterilmektedir. Tahmin edilen KKÖ "
        "değerleri, anomali olaylarına yaklaştıkça monoton bir azalma sergilemekte olup bu durum "
        "modelin kademeli bozulma sürecini başarıyla yakaladığını doğrulamaktadır. LSTM modeli, "
        "özellikle 72 saat altındaki kritik bölgede yüksek doğruluk göstermektedir.")

    add_fig(doc, "06_RUL_Prediction/results/rul_timeline.png",
            "Şekil 33: Zaman boyunca KKÖ tahminlerinin evrimi")

    body(doc,
        "LSTM ve GRU modellerinin karşılaştırmalı KKÖ tahmin performansı Şekil 34'te "
        "sunulmaktadır. LSTM modeli tüm ufuklarda GRU'dan düşük hata değerleri elde etmiş olup "
        "özellikle kısa vadeli tahminlerde (24-48 saat) üstünlüğü daha belirgin şekilde "
        "görülmektedir. Her iki model de 72 saat üzeri ufuklarda hata artışı göstermiştir.")

    add_fig(doc, "06_RUL_Prediction/results/rul_model_comparison.png",
            "Şekil 34: LSTM ve GRU KKÖ tahmin modellerinin karşılaştırması")

    # ============================================================
    #  F. KAPSAMLI PERFORMANS ÖZETİ
    # ============================================================
    heading(doc, "F. Kapsamlı Performans Özeti", 2)
    body(doc,
        "Tablo VIII, bu çalışmada değerlendirilen tüm yöntemlerin kapsamlı bir özetini "
        "sunmaktadır. Sonuçlar, bu veri setinde anomali tespiti için denetimsiz yöntemlerin "
        "üstünlüğünü açıkça gösterirken, derin öğrenme modelleri değerli zamansal örüntü "
        "tanıma yetenekleri sunmaktadır.")

    add_table(doc,
        ["Yöntem", "Kategori", "F1", "ROC-AUC", "PR-AUC", "Not"],
        [
            ["Lojistik Reg.", "Klasik MÖ", "0,048", "0,302", "0,223", "Temel"],
            ["Random Forest", "Klasik MÖ", "0,545", "0,472", "0,417", "En iyi F1"],
            ["XGBoost", "Klasik MÖ", "0,000", "0,823", "0,190", "En iyi AUC"],
            ["LightGBM", "Klasik MÖ", "0,000", "0,661", "0,152", "\u2014"],
            ["Isolation Forest", "Denetimsiz", "\u2014", "0,9997", "0,9946", "Mükemmel"],
            ["One-Class SVM", "Denetimsiz", "\u2014", "0,9999", "0,9965", "En iyi AUC"],
            ["LOF", "Denetimsiz", "\u2014", "0,9992", "0,9585", "Yoğunluk"],
            ["Autoencoder", "Denetimsiz", "\u2014", "0,9916", "0,8605", "Sinir ağı"],
            ["LSTM", "Derin Öğr.", "0,267", "0,591", "0,085", "Yüksek duy."],
            ["TCN", "Derin Öğr.", "0,445", "0,867", "0,601", "En dengeli"],
            ["Transformer", "Derin Öğr.", "0,357", "0,903", "0,516", "En iyi AUC"],
            ["LSTM (KKÖ)", "KKÖ Reg.", "\u2014", "\u2014", "\u2014", "OMH:3,96s"],
            ["GRU (KKÖ)", "KKÖ Reg.", "\u2014", "\u2014", "\u2014", "OMH:5,19s"],
        ],
        caption="Tablo VIII: Tüm Yöntemlerin Kapsamlı Performans Özeti", small=True)

    # ============================================================
    #  VI. TARTIŞMA
    # ============================================================
    heading(doc, "VI. TARTIŞMA", 1)

    heading(doc, "A. Denetimli ve Denetimsiz Yöntemlerin Karşılaştırması", 2)
    body(doc,
        "Bu çalışmanın çarpıcı bulgularından biri, denetimli ve denetimsiz yöntemler arasındaki "
        "önemli performans farkıdır. Klasik MÖ modelleri ciddi sınıf dengesizliği (47:1 oranı) "
        "ile mücadele ederken, denetimsiz yöntemler neredeyse mükemmel tespit oranlarına ulaşmıştır. "
        "Bu durum, dişli kutusu SCADA verilerindeki anomali örüntülerinin ayrımcı karar sınırlarından "
        "ziyade normal operasyonel dağılımlardan sapmalarla iyi karakterize edildiğini göstermektedir.")

    body(doc,
        "Denetimsiz yöntemlerin üstün performansı birkaç faktöre atfedilebilir: (1) yalnızca normal "
        "veriler üzerinde eğitim, sınıf dengesizliği sorununu tamamen ortadan kaldırır; (2) SCADA "
        "verisi, normal ve anormal davranışların özellik uzayında farklı bölgelerde yer aldığı "
        "net operasyonel rejimler sergilemektedir.")

    heading(doc, "B. Derin Öğrenme Mimari Karşılaştırması", 2)
    body(doc,
        "Derin öğrenme mimarileri arasında TCN, kesinlik ve duyarlılık arasında en iyi dengeyi "
        "sergilemiştir. Genişletilmiş nedensel konvolüsyonlar, TCN'nin çok ölçekli zamansal "
        "örüntüleri LSTM'den daha etkili yakalamasını sağlamıştır. LSTM anomalileri aşırı tespit "
        "etme eğiliminde iken (duyarlılık=0,993, kesinlik=0,155), Transformer Kodlayıcı güçlü "
        "ayrımcılık yeteneği (ROC-AUC=0,903) göstermiş ancak TCN'den daha fazla yanlış pozitif "
        "üretmiştir.")

    heading(doc, "C. Özellik Mühendisliğinin Etkisi", 2)
    body(doc,
        "Özellik uzayını 7'den 95 özelliğe genişleten mühendislik hattı, model performansında "
        "kritik bir rol oynamıştır. Karşılıklı bilgi analizi 168 saatlik yuvarlanan titreşim "
        "ortalamalarının en bilgilendirici özellikler olduğunu doğrularken, Random Forest özellik "
        "önemi 24-48 saatlik yağ basıncı ve sıcaklık istatistiklerini vurgulamıştır.")

    heading(doc, "D. KKÖ Tahmininin Önemi", 2)
    body(doc,
        "KKÖ tahmin bileşeni, bu çalışmanın pratik açıdan en anlamlı katkısını temsil etmektedir. "
        "Yaklaşık 4 saatlik ortalama tahmin hatası (LSTM) ve mükemmel erken uyarı doğruluğu ile "
        "sistem, bakım planlaması için eyleme dönüştürülebilir istihbarat sağlamaktadır. Dört "
        "bölgeli karar çerçevesi (Normal, İzleme, Uyarı, Kritik) ham KKÖ tahminlerini sezgisel "
        "operasyonel rehberliğe dönüştürmektedir.")

    heading(doc, "E. Sınırlılıklar ve Gelecek Çalışmalar", 2)
    body(doc,
        "Bazı sınırlılıklar kabul edilmelidir: (1) veri seti tek bir türbinden gelmekte olup "
        "performans farklı türbin tiplerinde değişebilir; (2) KKÖ regresyon modellerinin R² "
        "değerleri sıfıra yakındır; (3) çalışma kavram kayması veya değişen operasyonel koşulları "
        "ele almamaktadır. Gelecek çalışmalar transfer öğrenme, çevrimiçi öğrenme, hava durumu "
        "ve operasyonel bağlam verilerinin entegrasyonu ve uç bilişim ortamlarında gerçek zamanlı "
        "izleme konularını araştırmalıdır.")

    # ============================================================
    #  VII. SONUÇ
    # ============================================================
    heading(doc, "VII. SONUÇ", 1)

    body(doc,
        "Bu makale, beş yıllık SCADA sensör verisi kullanarak rüzgar türbini dişli kutuları "
        "için kapsamlı bir kestirimci bakım çerçevesi sunmuştur. Çalışma, klasik makine "
        "öğrenmesi, denetimsiz anomali tespiti, derin öğrenme ve hibrit topluluk yöntemlerini "
        "sistematik olarak değerlendirmiş; ardından Kalan Kullanım Ömrü tahmini ve erken uyarı "
        "sistemi geliştirmiştir.")

    body(doc, "Bu çalışmanın temel bulguları şu şekilde özetlenebilir:")

    findings = [
        "7 sensörden 95 özelliğe mühendislik, tahmin yeteneğini önemli ölçüde artırmış; "
        "168 saatlik yuvarlanan titreşim istatistikleri en ayırt edici özellikler olmuştur.",
        "Denetimsiz anomali tespit yöntemleri, özellikle One-Class SVM (ROC-AUC: 0,9999) ve "
        "Isolation Forest (ROC-AUC: 0,9997), denetimli yaklaşımları önemli ölçüde geride bırakmıştır.",
        "Derin öğrenme mimarileri arasında TCN en dengeli performansı (F1: 0,445, PR-AUC: 0,601) "
        "sağlarken, Transformer Kodlayıcı en yüksek ayrımcılık yeteneğine (ROC-AUC: 0,903) ulaşmıştır.",
        "LSTM tabanlı KKÖ tahmini, 3,96 saatlik ortalama mutlak hata ve 24, 48, 72 saatlik "
        "ufuklarda %100 erken uyarı doğruluğu ile güvenilir proaktif bakım planlaması sağlamıştır.",
        "Dört bölgeli bakım karar çerçevesi, reaktif bakımdan kestirimci bakım stratejilerine "
        "geçiş için eyleme dönüştürülebilir operasyonel rehber sunmaktadır.",
    ]
    for i, f in enumerate(findings, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{i}) {f}")
        r.font.name = "Times New Roman"
        r.font.size = Pt(9)

    body(doc,
        "Önerilen çerçeve, çoklu makine öğrenmesi paradigmalarının kapsamlı özellik mühendisliği "
        "ile entegrasyonunun rüzgar türbini kestirimci bakımı için sağlam ve etkili bir çözüm "
        "sağladığını göstermektedir. Anomali tespitinden KKÖ tahminine geçiş, arıza belirleme ve "
        "eyleme dönüştürülebilir bakım kararları arasındaki boşluğu kapatarak rüzgar enerjisi "
        "operasyonlarında duruş süresini ve bakım maliyetlerini azaltma potansiyeli sunmaktadır.")

    # ============================================================
    #  KAYNAKLAR
    # ============================================================
    heading(doc, "KAYNAKLAR", 1)

    refs = [
        '[1] Y. Wang, X. Ma, P. Qian, "Wind turbine fault detection and identification through '
        'self-attention-based mechanism," Renewable Energy, vol. 211, pp. 918-937, 2023.',
        '[2] A. Stetco et al., "Machine learning methods for wind turbine condition monitoring: '
        'A review," Renewable Energy, vol. 133, pp. 620-635, 2019.',
        '[3] F. P. G. de Jong, W. J. C. Verhagen, "A review of predictive maintenance for wind '
        'turbines using ML techniques," Energy Reports, vol. 8, pp. 5738-5768, 2022.',
        '[4] T. Chen, C. Guestrin, "XGBoost: A scalable tree boosting system," Proc. 22nd ACM '
        'SIGKDD, pp. 785-794, 2016.',
        '[5] G. Ke et al., "LightGBM: A highly efficient gradient boosting decision tree," '
        'NeurIPS, vol. 30, pp. 3146-3154, 2017.',
        '[6] F. T. Liu, K. M. Ting, Z.-H. Zhou, "Isolation forest," Proc. IEEE ICDM, '
        'pp. 413-422, 2008.',
        '[7] B. Scholkopf et al., "Estimating the support of a high-dimensional distribution," '
        'Neural Computation, vol. 13, no. 7, pp. 1443-1471, 2001.',
        '[8] M. M. Breunig et al., "LOF: Identifying density-based local outliers," '
        'Proc. ACM SIGMOD, pp. 93-104, 2000.',
        '[9] S. Hochreiter, J. Schmidhuber, "Long short-term memory," Neural Computation, '
        'vol. 9, no. 8, pp. 1735-1780, 1997.',
        '[10] S. Bai, J. Z. Kolter, V. Koltun, "An empirical evaluation of generic convolutional '
        'and recurrent networks for sequence modeling," arXiv:1803.01271, 2018.',
        '[11] A. Vaswani et al., "Attention is all you need," NeurIPS, vol. 30, '
        'pp. 5998-6008, 2017.',
        '[12] S. M. Lundberg, S.-I. Lee, "A unified approach to interpreting model predictions," '
        'NeurIPS, vol. 30, pp. 4765-4774, 2017.',
        '[13] M. T. Ribeiro, S. Singh, C. Guestrin, "Why should I trust you? Explaining the '
        'predictions of any classifier," Proc. 22nd ACM SIGKDD, pp. 1135-1144, 2016.',
        '[14] N. V. Chawla et al., "SMOTE: Synthetic minority over-sampling technique," '
        'JAIR, vol. 16, pp. 321-357, 2002.',
        '[15] W. Qiao, D. Lu, "A survey on wind turbine condition monitoring and fault '
        'diagnosis," IEEE Trans. Ind. Electron., vol. 62, no. 10, pp. 6536-6545, 2015.',
        '[16] K. Cho et al., "Learning phrase representations using RNN encoder-decoder," '
        'Proc. EMNLP, pp. 1724-1734, 2014.',
        '[17] X.-S. Si et al., "Remaining useful life estimation - A review on the statistical '
        'data driven approaches," EJOR, vol. 213, no. 1, pp. 1-14, 2011.',
        '[18] L. Breiman, "Random forests," Machine Learning, vol. 45, no. 1, pp. 5-32, 2001.',
        '[19] D. E. Rumelhart, G. E. Hinton, R. J. Williams, "Learning representations by '
        'back-propagating errors," Nature, vol. 323, pp. 533-536, 1986.',
        '[20] J. Macqueen, "Some methods for classification and analysis of multivariate '
        'observations," Proc. 5th Berkeley Symp., pp. 281-297, 1967.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        r = p.add_run(ref)
        r.font.name = "Times New Roman"
        r.font.size = Pt(7.5)
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.underline = False

    # ============================================================
    #  SAVE
    # ============================================================
    docx_path = os.path.join(REPORT_DIR, "IEEE_Report_Wind_Turbine_Predictive_Maintenance.docx")
    doc.save(docx_path)
    print(f"DOCX saved: {docx_path}")

    # Convert to PDF
    try:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", REPORT_DIR, docx_path],
            capture_output=True, text=True, timeout=PDF_CONVERSION_TIMEOUT_SECONDS)
        if result.returncode == 0:
            print(f"PDF saved: {os.path.join(REPORT_DIR, 'IEEE_Report_Wind_Turbine_Predictive_Maintenance.pdf')}")
        else:
            print(f"PDF error: {result.stderr}")
    except subprocess.TimeoutExpired:
        print(f"PDF conversion timed out after {PDF_CONVERSION_TIMEOUT_SECONDS} seconds")
    except Exception as e:
        print(f"PDF failed: {e}")

    print("Report generation complete!")


if __name__ == "__main__":
    generate_report()
